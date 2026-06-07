from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from django.shortcuts import get_object_or_404
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from calendar import monthrange
from datetime import date as dt_date, timedelta
from .models import (
    StatusReport, TrainingMaterial, ReportingItem, Meeting,
    GeneratedStatusReport, MethodologyReport,
)
from .serializers import (
    StatusReportSerializer, TrainingMaterialSerializer, ReportingItemSerializer,
    MeetingSerializer, MeetingOccurrenceSerializer, GeneratedStatusReportSerializer,
    MethodologyReportSerializer,
)


def _company_scoped(qs, user):
    """P0 cross-tenant fix — was returning every project's rows globally.

    All four communication models are scoped via `project.company`. Filter
    by the requesting user's company. Superadmin/superuser sees everything.
    """
    if not user.is_authenticated:
        return qs.none()
    if getattr(user, "role", None) == "superadmin" or getattr(user, "is_superuser", False):
        return qs
    company_id = getattr(user, "company_id", None)
    if not company_id:
        return qs.none()
    return qs.filter(project__company_id=company_id)


def _build_minutes_text(meeting) -> str:
    """Plain-text minutes for email."""
    lines = [f"MINUTES — {meeting.name}", ""]
    when = " ".join(filter(None, [str(meeting.date or ""), str(meeting.time or "")])).strip()
    if when:
        lines.append(f"Date/time: {when}")
    loc = meeting.location or getattr(meeting, "yanmar_meeting_room", "") or ""
    if loc:
        lines.append(f"Location: {loc}")
    atts = list(meeting.attendees.all())
    if atts:
        lines.append("")
        lines.append("Attendees:")
        for a in atts:
            tag = getattr(a, "presence", "") or ""
            lines.append(f"  - {a.name_text or ''} ({tag})".rstrip())
    agenda = meeting.agenda if isinstance(meeting.agenda, list) else []
    if agenda:
        lines += ["", "Agenda:"] + [f"  {i+1}. {a}" for i, a in enumerate(agenda)]
    if getattr(meeting, "discussion_notes", ""):
        lines += ["", "Discussion:", meeting.discussion_notes]
    if getattr(meeting, "conclusions", ""):
        lines += ["", "Conclusions:", meeting.conclusions]
    acts = list(meeting.action_items.all())
    if acts:
        lines += ["", "Action items:"]
        for it in acts:
            pic = it.pic_text or (it.pic_user.email if it.pic_user else "") or "—"
            lines.append(f"  - {it.subject} | PIC: {pic} | Due: {it.action_due or '—'} | {it.status}")
    return "\n".join(lines)


def _gather_recipients(meeting, extra) -> list:
    """Attendee emails (contact_info / name that looks like an email) + any extra
    addresses the caller passed (list or comma/newline-separated string)."""
    import re as _re
    out = set()
    for a in meeting.attendees.all():
        for cand in [getattr(a, "contact_info", ""), getattr(a, "name_text", "")]:
            if cand and _re.match(r"[^@\s]+@[^@\s]+\.[^@\s]+", cand.strip()):
                out.add(cand.strip())
    if isinstance(extra, str):
        extra = _re.split(r"[,\n;]+", extra)
    for e in (extra or []):
        e = (e or "").strip()
        if _re.match(r"[^@\s]+@[^@\s]+\.[^@\s]+", e):
            out.add(e)
    return sorted(out)


def _ics_escape(s: str) -> str:
    return (s or "").replace("\\", "\\\\").replace(";", "\\;").replace(",", "\\,").replace("\n", "\\n")


def _build_ics(meeting) -> str:
    """Minimal RFC-5545 VEVENT (all-day if no time, else a 1-hour slot)."""
    from datetime import datetime, timedelta
    uid = f"meeting-{meeting.id}@projextpal.com"
    summary = _ics_escape(meeting.name or "Meeting")
    loc = _ics_escape(meeting.location or getattr(meeting, "yanmar_meeting_room", "") or "")
    agenda = meeting.agenda if isinstance(meeting.agenda, list) else []
    desc = _ics_escape("Agenda:\n" + "\n".join(f"{i+1}. {a}" for i, a in enumerate(agenda)) if agenda else "")
    lines = ["BEGIN:VCALENDAR", "VERSION:2.0", "PRODID:-//ProjeXtPal//Meetings//EN", "BEGIN:VEVENT", f"UID:{uid}"]
    d = meeting.date
    if d and meeting.time:
        try:
            start = datetime.combine(d, meeting.time)
            end = start + timedelta(hours=1)
            lines += [f"DTSTART:{start.strftime('%Y%m%dT%H%M%S')}", f"DTEND:{end.strftime('%Y%m%dT%H%M%S')}"]
        except Exception:
            lines += [f"DTSTART;VALUE=DATE:{d.strftime('%Y%m%d')}"]
    elif d:
        lines += [f"DTSTART;VALUE=DATE:{d.strftime('%Y%m%d')}"]
    lines += [f"SUMMARY:{summary}"]
    if loc:
        lines += [f"LOCATION:{loc}"]
    if desc:
        lines += [f"DESCRIPTION:{desc}"]
    lines += ["END:VEVENT", "END:VCALENDAR"]
    return "\r\n".join(lines)


class StatusReportViewSet(viewsets.ModelViewSet):
    queryset = StatusReport.objects.all()
    serializer_class = StatusReportSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return _company_scoped(StatusReport.objects.all(), self.request.user)

    def list(self, request):
        queryset = self.get_queryset()
        project_id = request.query_params.get('project')
        if project_id:
            queryset = queryset.filter(project_id=project_id)
        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)
    
    def retrieve(self, request, pk=None):
        queryset = self.get_queryset()
        report = get_object_or_404(queryset, pk=pk)
        serializer = self.get_serializer(report)
        return Response(serializer.data)
    
    def create(self, request):
        serializer = self.get_serializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def update(self, request, pk=None):
        report = get_object_or_404(self.get_queryset(), pk=pk)
        serializer = self.get_serializer(report, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def partial_update(self, request, pk=None):
        report = get_object_or_404(self.get_queryset(), pk=pk)
        serializer = self.get_serializer(report, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def destroy(self, request, pk=None):
        report = get_object_or_404(self.get_queryset(), pk=pk)
        report.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


class GeneratedStatusReportViewSet(viewsets.ReadOnlyModelViewSet):
    """AI-synthesised executive status reports (IL-2).

    Read-only history + a `generate` action that runs the synthesis engine for a
    project and persists a new report. Company-scoped like the other surfaces.
    """
    serializer_class = GeneratedStatusReportSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = _company_scoped(
            GeneratedStatusReport.objects.select_related("project", "created_by"),
            self.request.user,
        )
        project_id = self.request.query_params.get("project")
        if project_id:
            qs = qs.filter(project_id=project_id)
        return qs

    @action(detail=False, methods=["post"])
    def generate(self, request):
        """Body: {project: <id>}. Gathers live metrics, computes RAG, synthesises
        the narrative (LLM if configured, else deterministic) and stores it."""
        from projects.models import Project
        from .status_synthesis import synthesize

        project_id = request.data.get("project")
        if not project_id:
            return Response({"detail": "project is required."},
                            status=status.HTTP_400_BAD_REQUEST)
        proj_qs = Project.objects.all()
        user = request.user
        if not (getattr(user, "role", None) == "superadmin" or getattr(user, "is_superuser", False)):
            company_id = getattr(user, "company_id", None)
            proj_qs = proj_qs.filter(company_id=company_id) if company_id else proj_qs.none()
        project = get_object_or_404(proj_qs, pk=project_id)
        result = synthesize(project)
        report = GeneratedStatusReport.objects.create(
            project=project,
            metrics=result["metrics"],
            overall_rag=result["overall_rag"],
            rag_scope=result["rag_scope"],
            rag_schedule=result["rag_schedule"],
            rag_cost=result["rag_cost"],
            rag_risk=result["rag_risk"],
            executive_summary=result["executive_summary"],
            highlights=result["highlights"],
            blockers=result["blockers"],
            next_steps=result["next_steps"],
            model_used=result["model_used"],
            original_ai_response=result["original_ai_response"],
            created_by=request.user,
        )
        return Response(self.get_serializer(report).data,
                        status=status.HTTP_201_CREATED)


class MethodologyReportViewSet(viewsets.ModelViewSet):
    """Doctrine-specific reports (one shared model, all methodologies).

    Full CRUD so a user can hand-author/edit a report, plus a `generate`
    action that synthesises one from the project's live methodology state.
    Company-scoped; filterable by `project` and `methodology`.
    """
    serializer_class = MethodologyReportSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = _company_scoped(
            MethodologyReport.objects.select_related("project", "created_by"),
            self.request.user,
        )
        project_id = self.request.query_params.get("project")
        if project_id:
            qs = qs.filter(project_id=project_id)
        methodology = self.request.query_params.get("methodology")
        if methodology:
            qs = qs.filter(methodology=methodology)
        return qs

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user, auto_generated=False)

    @action(detail=False, methods=["post"])
    def generate(self, request):
        """Body: {project: <id>, methodology: <slug>}. Reads live methodology
        state, synthesises the doctrine report and persists it."""
        from projects.models import Project
        from .doctrine_reports import synthesize

        project_id = request.data.get("project")
        methodology = request.data.get("methodology")
        if not project_id or not methodology:
            return Response({"detail": "project and methodology are required."},
                            status=status.HTTP_400_BAD_REQUEST)
        if methodology not in dict(MethodologyReport.METHODOLOGY_CHOICES):
            return Response({"detail": f"Unknown methodology '{methodology}'."},
                            status=status.HTTP_400_BAD_REQUEST)

        proj_qs = Project.objects.all()
        user = request.user
        if not (getattr(user, "role", None) == "superadmin" or getattr(user, "is_superuser", False)):
            company_id = getattr(user, "company_id", None)
            proj_qs = proj_qs.filter(company_id=company_id) if company_id else proj_qs.none()
        project = get_object_or_404(proj_qs, pk=project_id)

        result = synthesize(project, methodology)
        report = MethodologyReport.objects.create(
            project=project,
            methodology=methodology,
            report_type=result["report_type"],
            title=result["title"],
            period_start=result["period_start"],
            period_end=result["period_end"],
            scope_ref=result["scope_ref"],
            overall_rag=result["overall_rag"],
            rag_scope=result["rag_scope"],
            rag_schedule=result["rag_schedule"],
            rag_cost=result["rag_cost"],
            rag_risk=result["rag_risk"],
            executive_summary=result["executive_summary"],
            highlights=result["highlights"],
            blockers=result["blockers"],
            next_steps=result["next_steps"],
            metrics=result["metrics"],
            payload=result["payload"],
            auto_generated=True,
            model_used=result["model_used"],
            original_ai_response=result["original_ai_response"],
            created_by=request.user,
        )
        return Response(self.get_serializer(report).data,
                        status=status.HTTP_201_CREATED)


# Training Material ViewSet
class TrainingMaterialViewSet(viewsets.ModelViewSet):
    queryset = TrainingMaterial.objects.all()
    serializer_class = TrainingMaterialSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return _company_scoped(TrainingMaterial.objects.all(), self.request.user)

    def list(self, request):
        queryset = self.get_queryset()
        project_id = request.query_params.get('project')
        if project_id:
            queryset = queryset.filter(project_id=project_id)
        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)

    def retrieve(self, request, pk=None):
        queryset = self.get_queryset()
        material = get_object_or_404(queryset, pk=pk)
        serializer = self.get_serializer(material)
        return Response(serializer.data)

    def create(self, request):
        serializer = self.get_serializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def update(self, request, pk=None):
        material = get_object_or_404(self.get_queryset(), pk=pk)
        serializer = self.get_serializer(material, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def partial_update(self, request, pk=None):
        material = get_object_or_404(self.get_queryset(), pk=pk)
        serializer = self.get_serializer(material, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def destroy(self, request, pk=None):
        material = get_object_or_404(self.get_queryset(), pk=pk)
        material.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)
    
# Report ViewSet
class ReportingItemViewSet(viewsets.ModelViewSet):
    queryset = ReportingItem.objects.all()
    serializer_class = ReportingItemSerializer
    parser_classes = [MultiPartParser, FormParser, JSONParser]
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return _company_scoped(ReportingItem.objects.all(), self.request.user)

    def list(self, request, *args, **kwargs):
        qs = self.get_queryset()
        project_id = request.query_params.get('project')
        if project_id:
            qs = qs.filter(project_id=project_id)
        page = self.paginate_queryset(qs)
        serializer = self.get_serializer(page or qs, many=True, context={'request': request})
        if page is not None:
            return self.get_paginated_response(serializer.data)
        return Response(serializer.data)

    def retrieve(self, request, pk=None):
        item = get_object_or_404(self.get_queryset(), pk=pk)
        serializer = self.get_serializer(item, context={'request': request})
        return Response(serializer.data)

    def create(self, request):
        serializer = self.get_serializer(data=request.data, context={'request': request})
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    def partial_update(self, request, pk=None):
        item = get_object_or_404(self.get_queryset(), pk=pk)
        serializer = self.get_serializer(item, data=request.data, partial=True, context={'request': request})
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data)

    def destroy(self, request, pk=None):
        item = get_object_or_404(self.get_queryset(), pk=pk)
        item.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


# Meeting ViewSet
class MeetingViewSet(viewsets.ModelViewSet):
    queryset = Meeting.objects.all()
    serializer_class = MeetingSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = _company_scoped(Meeting.objects.all(), self.request.user)
        project_id = self.request.query_params.get('project')
        if project_id:
            qs = qs.filter(project_id=project_id)
        return qs

    @action(detail=True, methods=["post"], url_path="carry-forward")
    def carry_forward(self, request, pk=None):
        """MM-03 — clone OPEN action items from this meeting's previous_meeting
        into this meeting (skips subjects already present)."""
        meeting = self.get_object()
        prev = meeting.previous_meeting
        if not prev:
            # Not a client error — there is simply nothing to carry forward yet.
            return Response(
                {"carried_forward": 0, "detail": "No previous meeting set."},
                status=200,
            )
        existing = set(meeting.action_items.values_list("subject", flat=True))
        created = 0
        for item in prev.action_items.filter(status="open"):
            if item.subject in existing:
                continue
            item.carry_forward_to(meeting)
            created += 1
        return Response({"carried_forward": created}, status=200)

    @action(detail=True, methods=["post"], url_path="extract-minutes",
            parser_classes=[MultiPartParser, FormParser, JSONParser])
    def extract_minutes(self, request, pk=None):
        """Generate minutes + action items from raw notes — pasted `text` or an
        uploaded `file` (.docx/.txt). AI-extracts agenda/discussion/conclusions +
        actions and applies them to this meeting (actions appended, not replaced)."""
        from .minutes_extract import extract_text_from_upload, ai_extract_minutes
        from .models import MeetingActionItem

        meeting = self.get_object()
        text = (request.data.get("text") or "").strip()
        f = request.FILES.get("file")
        if f and not text:
            text = extract_text_from_upload(f).strip()
        if not text:
            return Response({"detail": "Provide notes text or a document.", "code": "no_input"}, status=400)

        parsed = ai_extract_minutes(text)
        if parsed.get("agenda"):
            meeting.agenda = parsed["agenda"]
        if parsed.get("discussion"):
            meeting.discussion_notes = parsed["discussion"]
        if parsed.get("conclusions"):
            meeting.conclusions = parsed["conclusions"]
        meeting.save()

        existing = set(meeting.action_items.values_list("subject", flat=True))
        created = 0
        for a in parsed.get("actions", []):
            subj = (a.get("subject") or "").strip()
            if not subj or subj in existing:
                continue
            MeetingActionItem.objects.create(
                meeting=meeting, subject=subj[:255],
                pic_text=(a.get("pic") or "")[:120],
                action_due=a.get("due") or None, status="open",
            )
            existing.add(subj)
            created += 1
        return Response({
            "applied": True, "actions_created": created,
            "agenda_count": len(parsed.get("agenda") or []),
            "meeting": MeetingSerializer(meeting).data,
        }, status=200)

    @action(detail=True, methods=["post"], url_path="push-actions-to-tasks")
    def push_actions_to_tasks(self, request, pk=None):
        """Push this meeting's OPEN action items into the project task list /
        Action Tracker — each becomes a Task (under a 'Meeting Actions' milestone)
        with owner (PIC) + due date, so meeting actions live in one activity list."""
        from datetime import datetime
        from projects.models import Task, Milestone

        meeting = self.get_object()
        project = meeting.project
        if not project:
            return Response({"detail": "Meeting has no project.", "code": "no_project"}, status=400)
        ms, _ = Milestone.objects.get_or_create(
            project=project, name="Meeting Actions",
            defaults={"description": "Actions captured from meetings", "status": "in_progress"},
        )
        existing = set(Task.objects.filter(milestone__project=project).values_list("title", flat=True))
        created = 0
        for it in meeting.action_items.filter(status="open"):
            title = (it.subject or "").strip()
            if not title or title in existing:
                continue
            due = None
            for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y"):
                try:
                    due = datetime.strptime((it.action_due or "").strip(), fmt).date(); break
                except (ValueError, TypeError):
                    continue
            Task.objects.create(
                milestone=ms, title=title[:255],
                description=f"From meeting: {meeting.name}",
                category="Meeting", status="todo", priority="medium",
                due_date=due, assigned_to=it.pic_user,
            )
            existing.add(title)
            created += 1
        return Response({"created": created}, status=200)

    @action(detail=True, methods=["post"], url_path="email-minutes")
    def email_minutes(self, request, pk=None):
        """Email the formatted minutes to the meeting attendees (those with an
        email) + any extra recipients in the body. Uses the configured backend."""
        from django.core.mail import EmailMessage
        from django.conf import settings as dj_settings

        meeting = self.get_object()
        recipients = _gather_recipients(meeting, request.data.get("recipients"))
        if not recipients:
            return Response({"detail": "No recipient email addresses. Add attendee emails or enter recipients.",
                             "code": "no_recipients"}, status=400)
        try:
            EmailMessage(
                subject=f"Minutes — {meeting.name}",
                body=_build_minutes_text(meeting),
                from_email=getattr(dj_settings, "DEFAULT_FROM_EMAIL", None) or getattr(dj_settings, "EMAIL_HOST_USER", None),
                to=recipients,
            ).send(fail_silently=False)
        except Exception as e:
            return Response({"detail": f"Email failed: {e}", "code": "send_failed"}, status=502)
        return Response({"sent_to": recipients}, status=200)

    @action(detail=True, methods=["post"], url_path="email-invite")
    def email_invite(self, request, pk=None):
        """Distribute the meeting as a calendar invite — emails the agenda + an
        attached .ics so recipients can add it to their calendar. Recipients =
        attendees with email + any extra addresses passed in `recipients`."""
        from django.core.mail import EmailMessage
        from django.conf import settings as dj_settings

        meeting = self.get_object()
        recipients = _gather_recipients(meeting, request.data.get("recipients"))
        if not recipients:
            return Response({"detail": "No recipient email addresses. Add attendee emails or enter recipients.",
                             "code": "no_recipients"}, status=400)
        when = " ".join(filter(None, [str(meeting.date or ""), str(meeting.time or "")])).strip()
        agenda = meeting.agenda if isinstance(meeting.agenda, list) else []
        body = f"You are invited to: {meeting.name}\n"
        if when:
            body += f"When: {when}\n"
        loc = meeting.location or getattr(meeting, "yanmar_meeting_room", "") or ""
        if loc:
            body += f"Where: {loc}\n"
        if agenda:
            body += "\nAgenda:\n" + "\n".join(f"  {i+1}. {a}" for i, a in enumerate(agenda))
        body += "\n\n(Open the attached .ics to add this to your calendar.)"
        try:
            msg = EmailMessage(
                subject=f"Invitation — {meeting.name}",
                body=body,
                from_email=getattr(dj_settings, "DEFAULT_FROM_EMAIL", None) or getattr(dj_settings, "EMAIL_HOST_USER", None),
                to=recipients,
            )
            msg.attach(f"meeting-{meeting.id}.ics", _build_ics(meeting), "text/calendar; method=REQUEST")
            msg.send(fail_silently=False)
        except Exception as e:
            return Response({"detail": f"Email failed: {e}", "code": "send_failed"}, status=502)
        return Response({"sent_to": recipients}, status=200)

    @action(detail=True, methods=["get"], url_path="invite.ics")
    def invite_ics(self, request, pk=None):
        """Download a calendar invite (.ics) for the meeting — date/time/location/
        agenda — to add to a calendar or forward as an agenda invite."""
        from django.http import HttpResponse
        meeting = self.get_object()
        ics = _build_ics(meeting)
        resp = HttpResponse(ics, content_type="text/calendar; charset=utf-8")
        resp["Content-Disposition"] = f'attachment; filename="meeting-{meeting.id}.ics"'
        return resp

    def list(self, request, *args, **kwargs):
        """
        GET /communication/meetings/?project=<id>
        GET /communication/meetings/?project=<id>&month=YYYY-MM&expand=true
        """
        expand = request.query_params.get('expand', 'false').lower() == 'true'
        month = request.query_params.get('month')

        if not expand or not month:
            qs = self.get_queryset()
            ser = self.get_serializer(qs, many=True)
            return Response(ser.data)

        # Expand recurring meetings for requested month
        try:
            year_str, mon_str = month.split('-')
            year, mon = int(year_str), int(mon_str)
        except Exception:
            return Response({"detail": "Invalid month format. Use YYYY-MM."}, status=400)

        first = dt_date(year, mon, 1)
        last_day = monthrange(year, mon)[1]
        last = dt_date(year, mon, last_day)

        def include(d: dt_date) -> bool:
            return first <= d <= last

        occurrences = []
        base_qs = self.get_queryset()

        for m in base_qs:
            base = m.date
            payload = MeetingSerializer(m, context={'request': request}).data

            if m.type == "onetime":
                if include(base):
                    occurrences.append({**payload, "date": base})
                continue

            if m.frequency == "adhoc":
                if include(base):
                    occurrences.append({**payload, "date": base})
                continue

            if m.frequency in ("weekly", "biweekly"):
                step = 7 if m.frequency == "weekly" else 14
                d = base
                while d < first:
                    d += timedelta(days=step)
                while d <= last:
                    occurrences.append({**payload, "date": d})
                    d += timedelta(days=step)

            elif m.frequency == "monthly":
                d = base
                while d < first:
                    y = d.year + (1 if d.month == 12 else 0)
                    mo = 1 if d.month == 12 else d.month + 1
                    mdays = monthrange(y, mo)[1]
                    d = dt_date(y, mo, min(d.day, mdays))

                while d <= last:
                    occurrences.append({**payload, "date": d})
                    y = d.year + (1 if d.month == 12 else 0)
                    mo = 1 if d.month == 12 else d.month + 1
                    mdays = monthrange(y, mo)[1]
                    d = dt_date(y, mo, min(d.day, mdays))

        data = MeetingOccurrenceSerializer(occurrences, many=True).data
        return Response(data)

# ── Yanmar Meeting Minutes sub-resources (MM-01 / MM-02 / MM-03) ─────────
from rest_framework.decorators import action
from .models import MeetingAttendee, MeetingActionItem
from .serializers import MeetingAttendeeSerializer, MeetingActionItemSerializer


def _meeting_scoped(qs, user):
    """Scope meeting sub-resources via meeting.project.company."""
    if not user.is_authenticated:
        return qs.none()
    if getattr(user, "role", None) == "superadmin" or getattr(user, "is_superuser", False):
        return qs
    company_id = getattr(user, "company_id", None)
    if not company_id:
        return qs.none()
    return qs.filter(meeting__project__company_id=company_id)


class MeetingAttendeeViewSet(viewsets.ModelViewSet):
    queryset = MeetingAttendee.objects.all()
    serializer_class = MeetingAttendeeSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = _meeting_scoped(MeetingAttendee.objects.all(), self.request.user)
        meeting_id = self.request.query_params.get("meeting")
        if meeting_id:
            qs = qs.filter(meeting_id=meeting_id)
        return qs


class MeetingActionItemViewSet(viewsets.ModelViewSet):
    queryset = MeetingActionItem.objects.all()
    serializer_class = MeetingActionItemSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = _meeting_scoped(MeetingActionItem.objects.all(), self.request.user)
        meeting_id = self.request.query_params.get("meeting")
        if meeting_id:
            qs = qs.filter(meeting_id=meeting_id)
        return qs


# ── Generic report export (DOCX / XLSX from a sections payload) ───────────
# One endpoint for ALL reports — the frontend sends the same `sections` it
# renders on screen, so any report gains Word/Excel export without a
# per-report backend renderer. PDF/MD/CSV/JSON are produced client-side.
import io as _io
from rest_framework.views import APIView as _APIView


class ReportExportView(_APIView):
    """POST /api/v1/communication/reports/export/
    Body: {"format": "docx"|"xlsx", "title": str,
           "sections": [{"heading": str, "rows": [[label, value], ...], "text": str}]}"""
    permission_classes = [IsAuthenticated]

    def post(self, request):
        fmt = (request.data.get("format") or "docx").lower()
        title = request.data.get("title") or "Report"
        sections = request.data.get("sections") or []
        if fmt == "xlsx":
            content, ext, ctype = self._xlsx(title, sections)
        else:
            content, ext, ctype = self._docx(title, sections)
        safe = "".join(c if c.isalnum() or c in "-_" else "-" for c in title)[:60] or "report"
        from django.http import HttpResponse
        r = HttpResponse(content, content_type=ctype)
        r["Content-Disposition"] = f'attachment; filename="{safe}.{ext}"'
        return r

    def _docx(self, title, sections):
        from docx import Document
        doc = Document()
        doc.add_heading(title, level=0)
        for s in sections:
            if s.get("heading"):
                doc.add_heading(str(s["heading"]), level=1)
            rows = s.get("rows") or []
            if rows:
                table = doc.add_table(rows=0, cols=2)
                table.style = "Light Grid Accent 1"
                for pair in rows:
                    label = str(pair[0]) if len(pair) > 0 else ""
                    value = str(pair[1]) if len(pair) > 1 else ""
                    cells = table.add_row().cells
                    cells[0].text = label
                    cells[1].text = value
            if s.get("text"):
                doc.add_paragraph(str(s["text"]))
        buf = _io.BytesIO()
        doc.save(buf)
        return buf.getvalue(), "docx", (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def _xlsx(self, title, sections):
        from openpyxl import Workbook
        from openpyxl.styles import Font
        wb = Workbook()
        ws = wb.active
        ws.title = "Report"
        ws.append([title])
        ws["A1"].font = Font(bold=True, size=14)
        ws.append([])
        for s in sections:
            if s.get("heading"):
                ws.append([str(s["heading"])])
                ws.cell(row=ws.max_row, column=1).font = Font(bold=True)
            for pair in (s.get("rows") or []):
                label = str(pair[0]) if len(pair) > 0 else ""
                value = str(pair[1]) if len(pair) > 1 else ""
                ws.append([label, value])
            if s.get("text"):
                ws.append(["", str(s["text"])])
            ws.append([])
        ws.column_dimensions["A"].width = 28
        ws.column_dimensions["B"].width = 60
        buf = _io.BytesIO()
        wb.save(buf)
        return buf.getvalue(), "xlsx", (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
