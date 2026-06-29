"""
AI Meeting Minutes generator.

Takes a raw transcript (text), summarises it via Claude into the structured
fields used by the Yanmar Meeting Minutes template, and renders a DOCX that
mirrors the template's layout exactly:

    Header table        - Date / Time / Prepared by / Location
                          Customer-Supplier | Yanmar Europe meeting rooms
    Attendees table     - Invited / Attendees / Absent (Name / Position / Contact)
    Reason(s) for the meeting
    Agenda points
    Previous Actions    - (carried in from the caller)
    Discussion Points
    Agreed New Actions
    Conclusions

The Claude prompt is deliberately structured so the response is JSON we can
load directly into the renderer. If the API call fails or returns invalid
JSON, the caller receives an error -- we never write a half-rendered DOCX.
"""

from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass, field
from datetime import date, datetime
from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

# Anthropic SDK is optional at import time so the rest of the app loads
# even if the package isn't installed in dev.
try:
    from anthropic import Anthropic
    _ANTHROPIC_OK = True
except Exception:  # pragma: no cover
    Anthropic = None  # type: ignore[assignment]
    _ANTHROPIC_OK = False


# ---------------------------------------------------------------------------
# DTOs
# ---------------------------------------------------------------------------

@dataclass
class Attendee:
    name: str = ""
    position: str = ""
    contact: str = ""
    presence: str = "attended"  # invited | attended | absent


@dataclass
class ActionItem:
    no: int = 0
    subject: str = ""
    pic: str = ""           # Person In Charge
    action_due: str = ""    # free-text date


@dataclass
class MinutesData:
    meeting_name: str = ""
    meeting_date: Optional[date] = None
    start_time: str = ""
    end_time: str = ""
    location: str = "Online Meeting"
    prepared_by: str = ""
    customer_supplier: str = ""
    yanmar_meeting_room: str = ""

    attendees: list[Attendee] = field(default_factory=list)
    reason: str = ""
    agenda_points: list[str] = field(default_factory=list)
    previous_actions: list[ActionItem] = field(default_factory=list)
    discussion_points: list[str] = field(default_factory=list)
    new_actions: list[ActionItem] = field(default_factory=list)
    conclusions: str = ""


# ---------------------------------------------------------------------------
# Claude prompt
# ---------------------------------------------------------------------------

SYSTEM_PROMPT = """You are an expert minute-taker for Yanmar Europe project meetings.
Given a raw meeting transcript, produce structured meeting minutes in the
Yanmar template format. You MUST output a single JSON object and nothing else.

The JSON schema:
{
  "meeting_name": str,
  "meeting_date": "YYYY-MM-DD" | null,
  "start_time": "HH:MM" | "",
  "end_time": "HH:MM" | "",
  "location": str,
  "prepared_by": str,
  "customer_supplier": str,
  "yanmar_meeting_room": str,
  "attendees": [
    {"name": str, "position": str, "contact": str,
     "presence": "invited" | "attended" | "absent"}
  ],
  "reason": str,
  "agenda_points": [str],
  "discussion_points": [str],   // 3-8 concise bullets
  "new_actions": [
    {"no": int, "subject": str, "pic": str, "action_due": str}
  ],
  "conclusions": str
}

Rules:
- Be faithful to the transcript; do not invent attendees or actions.
- Action items must each have a clear subject, an explicit PIC, and a due date
  if mentioned. If a due date is implied ("next week", "before Friday"),
  preserve the original phrasing.
- "discussion_points" should be a compact bulleted summary, not a paraphrase
  of every line.
- Use Yanmar Europe vocabulary: "PIC" for owner, "agenda points" not "agenda".
- Output JSON only -- no commentary, no markdown fence.
"""


def _strip_code_fence(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        if text.endswith("```"):
            text = text[:-3]
    return text.strip()


def transcript_to_minutes(
    transcript: str,
    *,
    api_key: Optional[str] = None,
    base_url: Optional[str] = None,
    project_name: str = "",
    previous_actions: Optional[list[ActionItem]] = None,
    model: str = "claude-opus-4-7",
) -> MinutesData:
    """
    Send transcript to Claude, parse JSON, return a fully-populated MinutesData.
    Raises RuntimeError on missing API key or invalid response.

    `base_url` is optional and lets a company route Anthropic traffic via a
    proxy (AWS Bedrock gateway, enterprise egress, etc.).
    """
    if not _ANTHROPIC_OK:
        raise RuntimeError(
            "anthropic SDK not installed -- add 'anthropic' to requirements.txt"
        )
    key = api_key or os.environ.get("ANTHROPIC_API_KEY") or ""
    if not key:
        raise RuntimeError("ANTHROPIC_API_KEY not configured")

    client_kwargs: dict = {"api_key": key}
    if base_url:
        client_kwargs["base_url"] = base_url
    client = Anthropic(**client_kwargs)
    user_msg = (
        f"Project context: {project_name or 'unspecified'}\n\n"
        "Transcript:\n\n"
        f"{transcript.strip()}"
    )
    resp = client.messages.create(
        model=model,
        max_tokens=4000,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_msg}],
    )
    # Concatenate all text blocks (Claude sometimes splits)
    raw = "".join(
        block.text for block in resp.content if getattr(block, "type", "") == "text"
    )
    raw = _strip_code_fence(raw)
    try:
        payload = json.loads(raw)
    except json.JSONDecodeError as e:
        raise RuntimeError(f"Claude did not return valid JSON: {e}\nRaw: {raw[:300]}") from e

    md = payload.get("meeting_date")
    meeting_date = None
    if md:
        try:
            meeting_date = datetime.strptime(md, "%Y-%m-%d").date()
        except ValueError:
            meeting_date = None

    return MinutesData(
        meeting_name=payload.get("meeting_name", ""),
        meeting_date=meeting_date,
        start_time=payload.get("start_time", "") or "",
        end_time=payload.get("end_time", "") or "",
        location=payload.get("location", "Online Meeting"),
        prepared_by=payload.get("prepared_by", ""),
        customer_supplier=payload.get("customer_supplier", ""),
        yanmar_meeting_room=payload.get("yanmar_meeting_room", ""),
        attendees=[
            Attendee(
                name=a.get("name", ""),
                position=a.get("position", ""),
                contact=a.get("contact", ""),
                presence=a.get("presence", "attended"),
            )
            for a in payload.get("attendees", [])
        ],
        reason=payload.get("reason", ""),
        agenda_points=list(payload.get("agenda_points", [])),
        previous_actions=previous_actions or [],
        discussion_points=list(payload.get("discussion_points", [])),
        new_actions=[
            ActionItem(
                no=int(act.get("no", i + 1)),
                subject=act.get("subject", ""),
                pic=act.get("pic", ""),
                action_due=act.get("action_due", ""),
            )
            for i, act in enumerate(payload.get("new_actions", []))
        ],
        conclusions=payload.get("conclusions", ""),
    )


# ---------------------------------------------------------------------------
# DOCX renderer matching the Yanmar Minutes template
# ---------------------------------------------------------------------------

def _h(doc: Document, text: str, *, size: int = 11, bold: bool = True) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def _bullet(doc: Document, items: list[str]) -> None:
    if not items:
        p = doc.add_paragraph()
        p.add_run("—").italic = True
        return
    for it in items:
        if not it:
            continue
        doc.add_paragraph(it, style="List Bullet")


def _action_table(doc: Document, actions: list[ActionItem]) -> None:
    headers = ["No", "Subject", "Person In Charge (PIC)", "Action Due"]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for a in actions or []:
        row = table.add_row().cells
        row[0].text = str(a.no or "")
        row[1].text = a.subject or ""
        row[2].text = a.pic or ""
        row[3].text = a.action_due or ""
    # Ensure at least 5 visible rows like the template
    while len(table.rows) < 6:
        row = table.add_row().cells
        row[0].text = str(len(table.rows) - 1)


def render_minutes_docx(data: MinutesData) -> bytes:
    """Render MinutesData into the Yanmar Meeting Minutes DOCX layout."""
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Minutes of Meeting")
    run.bold = True
    run.font.size = Pt(18)

    name = doc.add_paragraph()
    run = name.add_run(f"Meeting Name: {data.meeting_name or 'TBD'}")
    run.bold = True
    run.font.size = Pt(12)

    # Header table (Date/Time/Prepared by + Location/Customer/Yanmar room)
    table = doc.add_table(rows=3, cols=6)
    table.style = "Light Grid Accent 1"
    date_str = data.meeting_date.isoformat() if data.meeting_date else ""
    time_str = (
        f"{data.start_time} / {data.end_time}"
        if data.start_time or data.end_time else "00:00 / 00:00"
    )
    rows = [
        ["Date and Time", f"{date_str} {time_str}".strip(),
         "Prepared by", data.prepared_by or "", "", ""],
        ["Meeting Location", data.location or "Online Meeting",
         "Customer/Supplier", data.customer_supplier or "",
         "Yanmar Europe", ""],
        ["", "", "", "",
         "Meeting room:", data.yanmar_meeting_room or ""],
    ]
    for ri, vals in enumerate(rows):
        for ci, v in enumerate(vals):
            cell = table.rows[ri].cells[ci]
            cell.text = v
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    if ci % 2 == 0 and v:
                        r.bold = True

    doc.add_paragraph()

    # Attendees
    _h(doc, "Attendees")
    att_table = doc.add_table(rows=1, cols=4)
    att_table.style = "Light Grid Accent 1"
    hdr = att_table.rows[0].cells
    for i, h in enumerate(["", "Name", "Position", "Contact Information"]):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    def _section(label: str, presence: str) -> None:
        marker = att_table.add_row().cells
        marker[0].text = label
        for r in marker[0].paragraphs[0].runs:
            r.bold = True
        any_in_section = False
        for a in data.attendees:
            if a.presence == presence:
                any_in_section = True
                row = att_table.add_row().cells
                row[0].text = ""
                row[1].text = a.name
                row[2].text = a.position
                row[3].text = a.contact
        if not any_in_section:
            row = att_table.add_row().cells

    _section("Invited", "invited")
    _section("Attendees", "attended")
    _section("Absent", "absent")

    doc.add_paragraph()
    _h(doc, "Reason(s) for the Meeting")
    doc.add_paragraph(data.reason or "—")

    doc.add_paragraph()
    _h(doc, "Agenda Points")
    _bullet(doc, data.agenda_points)

    doc.add_paragraph()
    _h(doc, "Previous Actions")
    _action_table(doc, data.previous_actions)

    doc.add_paragraph()
    _h(doc, "Discussion Points")
    _bullet(doc, data.discussion_points)

    doc.add_paragraph()
    _h(doc, "Agreed New Actions")
    _action_table(doc, data.new_actions)

    doc.add_paragraph()
    _h(doc, "Conclusions")
    doc.add_paragraph(data.conclusions or "—")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
