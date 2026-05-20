"""
DOCX export for the Yanmar "Project Plan" template.

Generates a Word document that matches the layout of
    `Sales/Yanmar/Project Plan - Template.docx`

Structure (matches Yanmar's template):
    Title
    Project Name / Team
    Roles table (Owner / PM / Leader / Facilitator / Outside Eyes / Stakeholders)
    Target / Problem Statement
    Scope (in + out)
    Impact / Solution
    Resources (people / cost / materials)
    Time frame (Start / Target Implementation / Target End)
    Action Plan table (No / Topic / Responsible / Deadline)
    Push-back rule notice
    Risks and Issues
    Communication plan
    Project Go-Live
    Project closing
"""

from __future__ import annotations

from io import BytesIO
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(rows):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = label
        c1.text = value or ""
        for run in c0.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)


def _date(d) -> str:
    return d.isoformat() if d else ""


def render_project_plan_docx(project) -> bytes:
    """Render a Yanmar Project Plan DOCX from a projects.Project instance."""
    doc = Document()

    _h1(doc, "Project Plan")

    # Project name + team
    _h2(doc, f"Project Name: {project.name}")
    doc.add_paragraph(" ")  # spacer

    # Roles -- prefer ProjectMembership (Yanmar's canonical 6-role taxonomy);
    # fall back to PRINCE2 ProjectBoardMember for projects that pre-date the
    # ProjectMembership model.
    _h2(doc, "Project Team")
    role_index = {
        "project_owner":   0,
        "project_manager": 1,
        "project_leader":  2,
        "facilitator":     3,
        "outside_eyes":    4,
        "stakeholder":     5,
    }
    role_rows = [
        ("Project Owner", ""),
        ("Project Manager", ""),
        ("Project Leader", ""),
        ("Facilitator", ""),
        ("Outside Eyes", ""),
        ("Stakeholder(s)", ""),
    ]

    # Pass 1 — unified ProjectMembership rows.
    try:
        for m in project.memberships.select_related("user").all():
            slot = role_index.get(m.role)
            if slot is None:
                continue
            user = m.user
            name = getattr(user, "first_name", "") or getattr(user, "email", "")
            existing = role_rows[slot][1]
            if existing:
                # Append additional people on the same role.
                role_rows[slot] = (role_rows[slot][0], f"{existing}, {name}")
            else:
                role_rows[slot] = (role_rows[slot][0], name)
    except Exception:
        pass

    # Pass 2 — legacy PRINCE2 fallback (only fills empty slots).
    try:
        board = getattr(project, "prince2_board", None)
        if board:
            for bm in board.members.all():
                role = (bm.role or "").lower()
                if bm.user_id:
                    user = bm.user
                    name = (
                        getattr(user, "first_name", "") or
                        getattr(user, "email", "")
                    )
                else:
                    name = ""
                mapping = {
                    "executive":       0,  # Project Owner
                    "project_manager": 1,
                    "senior_supplier": 4,  # Outside Eyes
                }
                slot = mapping.get(role)
                if slot is not None and not role_rows[slot][1]:
                    role_rows[slot] = (role_rows[slot][0], name)
    except Exception:
        pass

    _kv_table(doc, role_rows)

    doc.add_paragraph(" ")

    # Target / problem
    _h2(doc, "Target / Problem Statement")
    doc.add_paragraph(project.project_goal or project.description or "—")

    # Impact + Solution (Yanmar §3-4 — problem-statement projects)
    impact = getattr(project, "problem_impact", "") or ""
    solution = getattr(project, "proposed_solution", "") or ""
    if impact or solution:
        doc.add_paragraph(" ")
        _h2(doc, "Impact (if problem statement)")
        doc.add_paragraph(impact or "—")
        _h2(doc, "Solution (if problem statement)")
        doc.add_paragraph(solution or "—")

    # Scope
    doc.add_paragraph(" ")
    _h2(doc, "Scope (in)")
    doc.add_paragraph(getattr(project, "scope_in", "") or "—")
    _h2(doc, "Scope (out)")
    doc.add_paragraph(getattr(project, "scope_out", "") or "—")

    # ROI (only if set)
    roi_t = getattr(project, "roi_target_pct", None)
    roi_r = getattr(project, "roi_realized_pct", None)
    if roi_t is not None or roi_r is not None:
        doc.add_paragraph(" ")
        _h2(doc, "ROI")
        _kv_table(doc, [
            ("Target ROI %", f"{roi_t}%" if roi_t is not None else "—"),
            ("Realized ROI %", f"{roi_r}%" if roi_r is not None else "—"),
        ])

    # Resources
    doc.add_paragraph(" ")
    _h2(doc, "Resources")
    p = doc.add_paragraph()
    p.add_run("People (time): ").bold = True
    try:
        team_count = project.team_members.count()
    except Exception:
        team_count = "—"
    p.add_run(f"{team_count} team members\n")
    p.add_run("Cost overview (budget): ").bold = True
    budget = getattr(project, "budget", "") or "—"
    p.add_run(f"{budget} {getattr(project, 'currency', 'EUR')}\n")
    p.add_run("Materials (equipment, software, etc.): ").bold = True
    p.add_run("—\n")

    # Time frame
    doc.add_paragraph(" ")
    _h2(doc, "Time frame")
    _kv_table(doc, [
        ("Start date", _date(project.start_date)),
        ("Target Implementation date",
         _date(getattr(project, "target_implementation_date", None))),
        ("Target end date", _date(project.end_date)),
    ])

    # Action plan
    doc.add_paragraph(" ")
    _h2(doc, "Action Plan")
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["No", "Topic", "Responsible", "Deadline"]):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
    # Pull top-level milestones as action plan items (Yanmar template uses
    # a high-level list, not the full task list).
    n = 1
    try:
        for m in project.milestones.all().order_by("order_index", "id")[:10]:
            row = tbl.add_row().cells
            row[0].text = str(n)
            row[1].text = m.name or ""
            row[2].text = ""
            row[3].text = _date(getattr(m, "end_date", None))
            n += 1
    except Exception:
        pass
    while len(tbl.rows) < 6:
        row = tbl.add_row().cells
        row[0].text = str(n)
        n += 1

    # Push-back rule notice (literally from Yanmar's template).
    doc.add_paragraph(" ")
    note = doc.add_paragraph()
    note_run = note.add_run(
        "* Due dates can only be pushed back once for a maximum period of "
        "2 weeks. More than this should be discussed with and approved by "
        "the Project Owner."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)

    # Risks & Issues
    doc.add_paragraph(" ")
    _h2(doc, "Risks and Issues")
    try:
        risks = list(project.risks.all().values_list("name", flat=True))
    except Exception:
        risks = []
    if risks:
        for r in risks:
            doc.add_paragraph(r, style="List Bullet")
    else:
        doc.add_paragraph("—")

    # Communication plan
    doc.add_paragraph(" ")
    _h2(doc, "Communication plan")
    for line in [
        "Kick-off meeting (Initial planning, define who is involved)",
        "Onboarding of stakeholders meeting (Inform project team and stakeholders their individual tasks)",
        "Regular update meetings (Weekly / Bi-weekly / Monthly)",
        "Project completion meeting",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # Go-Live
    doc.add_paragraph(" ")
    _h2(doc, "Project Go-Live")
    doc.add_paragraph(
        f"Target implementation: "
        f"{_date(getattr(project, 'target_implementation_date', None)) or '—'}"
    )

    # Closing
    doc.add_paragraph(" ")
    _h2(doc, "Project closing")
    doc.add_paragraph(
        "2–4 weeks after completion, project team performs evaluation. "
        "Senior Manager signs off every project as a 'finished project'."
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
